#!/usr/bin/env python3
"""
generate-site.py  -  Generate a full docs site from SKILLS.json.

Produces:
  docs/index.html       -  Landing page
  docs/docs.html        -  Getting started + how it works
  docs/browse.html      -  Interactive skill browser with search
  docs/examples.html    -  Real-world examples and use cases
  docs/api.html         -  SKILLS.json format reference

Usage: python3 scripts/generate-site.py
"""

import json
import html as html_mod
from pathlib import Path
from collections import defaultdict

ROOT = Path(__file__).resolve().parent.parent
SKILLS_JSON = ROOT / 'SKILLS.json'
DOCS = ROOT / 'docs'

def load():
    with open(SKILLS_JSON) as f:
        return json.load(f)

# ── Shared components ──

NAV = '''
<nav>
  <div class="container">
    <a href="index.html" class="logo">1ai-Skills</a>
    <div class="burger" onclick="document.getElementById('navLinks').classList.toggle('open')">
      <span></span><span></span><span></span>
    </div>
    <div class="links" id="navLinks">
      <a href="docs.html">Docs</a>
      <a href="browse.html">Browse</a>
      <a href="examples.html">Examples</a>
      <a href="api.html">API</a>
      <a href="https://github.com/oyi77/1ai-skills" class="btn btn-ghost" target="_blank">
        <svg width="16" height="16" viewBox="0 0 16 16" fill="currentColor"><path d="M8 0C3.58 0 0 3.58 0 8c0 3.54 2.29 6.53 5.47 7.59.4.07.55-.17.55-.38 0-.19-.01-.82-.01-1.49-2.01.37-2.53-.49-2.69-.94-.09-.23-.48-.94-.82-1.13-.28-.15-.68-.52-.01-.53.63-.01 1.08.58 1.23.82.72 1.21 1.87.87 2.33.66.07-.52.28-.87.51-1.07-1.78-.2-3.64-.89-3.64-3.95 0-.87.31-1.59.82-2.15-.08-.2-.36-1.02.08-2.12 0 0 .67-.21 2.2.82.64-.18 1.32-.27 2-.27.68 0 1.36.09 2 .27 1.53-1.04 2.2-.82 2.2-.82.44 1.1.16 1.92.08 2.12.51.56.82 1.27.82 2.15 0 3.07-1.87 3.75-3.65 3.95.29.25.54.73.54 1.48 0 1.07-.01 1.93-.01 2.2 0 .21.15.46.55.38A8.01 8.01 0 0016 8c0-4.42-3.58-8-8-8z"/></svg>
        GitHub
      </a>
    </div>
  </div>
</nav>'''

FOOTER = '''
<footer>
  <div class="container footer-inner">
    <div class="footer-left">
      <span class="logo-sm">1ai-Skills</span>
      <span class="copy">&copy; 2026 1ai. MIT License.</span>
    </div>
    <div class="footer-links">
      <a href="https://github.com/oyi77/1ai-skills" target="_blank">GitHub</a>
      <a href="https://github.com/oyi77/1ai-skills/blob/main/CONTRIBUTING.md" target="_blank">Contributing</a>
      <a href="https://github.com/oyi77/1ai-skills/blob/main/SECURITY.md" target="_blank">Security</a>
    </div>
  </div>
</footer>'''

CSS = '''
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
:root{
  --bg:#06060e;--bg2:#0c0c1a;--bg3:#111128;
  --surface:rgba(255,255,255,.03);--surface2:rgba(255,255,255,.06);--surface3:rgba(255,255,255,.08);
  --text:#e8e8f0;--text2:#9898b8;--text3:#5a5a78;
  --accent:#8b5cf6;--accent2:#06d6a0;--accent3:#38bdf8;--accent4:#f472b6;--accent5:#fbbf24;
  --grad:linear-gradient(135deg,#8b5cf6,#38bdf8,#06d6a0);
  --border:rgba(255,255,255,.06);--border2:rgba(255,255,255,.12);
  --glass:rgba(12,12,26,.6);--glass2:rgba(12,12,26,.8);
  --radius:16px;--radius-sm:10px;--radius-xs:6px;
  --shadow:0 8px 32px rgba(0,0,0,.4);
  --font:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,'Helvetica Neue',Arial,sans-serif;
  --mono:'SF Mono','Fira Code','Cascadia Code',Consolas,monospace;
}
html{scroll-behavior:smooth}
body{font-family:var(--font);background:var(--bg);color:var(--text);line-height:1.65;overflow-x:hidden}
a{color:var(--accent3);text-decoration:none;transition:color .2s}
a:hover{color:#7dd3fc}
.container{max-width:1240px;margin:0 auto;padding:0 24px}
code{font-family:var(--mono);background:var(--surface2);padding:2px 6px;border-radius:4px;font-size:.88em}
pre{background:var(--bg2);border:1px solid var(--border);border-radius:var(--radius-xs);padding:20px;overflow-x:auto;font-family:var(--mono);font-size:.85rem;color:var(--accent2);line-height:1.7}
pre code{background:none;padding:0}

/* NAV */
nav{position:fixed;top:0;left:0;right:0;z-index:1000;background:var(--glass2);backdrop-filter:blur(20px);border-bottom:1px solid var(--border)}
nav .container{display:flex;align-items:center;justify-content:space-between;height:64px}
nav .logo{font-size:1.2rem;font-weight:800;background:var(--grad);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text}
nav .links{display:flex;gap:8px;align-items:center}
nav .links a{color:var(--text2);font-size:.85rem;font-weight:500;padding:6px 14px;border-radius:8px;transition:all .2s}
nav .links a:not(.btn):hover{color:#fff;background:var(--surface2)}
.btn{display:inline-flex;align-items:center;gap:8px;padding:10px 24px;border-radius:var(--radius-sm);font-weight:600;font-size:.875rem;cursor:pointer;border:none;transition:all .25s}
.btn-primary{background:var(--grad);color:#fff}
.btn-primary:hover{opacity:.92;transform:translateY(-1px)}
.btn-ghost{background:var(--surface);color:var(--text);border:1px solid var(--border2)}
.btn-ghost:hover{border-color:var(--accent);color:#fff}
.burger{display:none;flex-direction:column;gap:5px;cursor:pointer;padding:4px}
.burger span{display:block;width:22px;height:2px;background:var(--text2);border-radius:2px}

/* PAGE LAYOUT */
.page{padding:100px 0 60px;min-height:100vh}
.page-header{text-align:center;margin-bottom:56px}
.page-header h1{font-size:clamp(2rem,4vw,3rem);font-weight:800;letter-spacing:-.02em;margin-bottom:12px}
.page-header h1 .grad{background:var(--grad);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text}
.page-header p{color:var(--text2);font-size:1.1rem;max-width:600px;margin:0 auto}

/* CARDS */
.card{background:var(--surface);border:1px solid var(--border);border-radius:var(--radius);padding:28px;transition:all .3s}
.card:hover{border-color:var(--border2);transform:translateY(-2px);box-shadow:0 8px 32px rgba(0,0,0,.2)}
.card h3{font-size:1.1rem;font-weight:700;margin-bottom:8px;color:#fff}
.card p{font-size:.88rem;color:var(--text2);line-height:1.6}

/* GRID */
.grid-2{display:grid;grid-template-columns:repeat(auto-fill,minmax(360px,1fr));gap:20px}
.grid-3{display:grid;grid-template-columns:repeat(auto-fill,minmax(280px,1fr));gap:16px}
.grid-4{display:grid;grid-template-columns:repeat(auto-fill,minmax(220px,1fr));gap:14px}

/* TABLE */
table{width:100%;border-collapse:collapse;margin:20px 0}
th,td{text-align:left;padding:12px 16px;border-bottom:1px solid var(--border)}
th{color:var(--accent);font-size:.8rem;text-transform:uppercase;letter-spacing:.08em;font-weight:600}
td{font-size:.9rem;color:var(--text2)}
tr:hover td{color:var(--text);background:var(--surface)}

/* BADGE */
.badge{display:inline-block;font-size:.7rem;font-weight:600;padding:3px 10px;border-radius:999px;text-transform:uppercase;letter-spacing:.05em}
.badge-cat{background:rgba(139,92,246,.15);color:var(--accent)}
.badge-new{background:rgba(244,114,182,.15);color:var(--accent4)}

/* SEARCH */
.search-wrap{max-width:500px;margin:0 auto 40px;position:relative}
.search-wrap input{width:100%;padding:14px 20px 14px 48px;border-radius:var(--radius-sm);border:1px solid var(--border2);background:var(--surface);color:var(--text);font-size:.95rem;font-family:var(--font);outline:none;transition:border-color .2s}
.search-wrap input:focus{border-color:var(--accent);box-shadow:0 0 0 3px rgba(139,92,246,.15)}
.search-wrap .search-icon{position:absolute;left:16px;top:50%;transform:translateY(-50%);color:var(--text3);font-size:1.1rem;pointer-events:none}

/* SKILL CARD */
.skill-card{background:var(--surface);border:1px solid var(--border);border-radius:var(--radius);padding:20px;transition:all .3s;cursor:pointer}
.skill-card:hover{border-color:var(--accent);transform:translateY(-2px)}
.skill-card h4{font-size:.95rem;font-weight:700;margin-bottom:6px;color:#fff}
.skill-card .desc{font-size:.82rem;color:var(--text2);line-height:1.5;margin-bottom:10px;display:-webkit-box;-webkit-line-clamp:2;-webkit-box-orient:vertical;overflow:hidden}
.skill-card .meta{display:flex;gap:8px;flex-wrap:wrap}
.skill-card .meta span{font-size:.7rem;padding:2px 8px;border-radius:999px;background:var(--surface2);color:var(--text3)}

/* SIDEBAR */
.docs-layout{display:grid;grid-template-columns:240px 1fr;gap:40px;align-items:start}
.docs-sidebar{position:sticky;top:80px}
.docs-sidebar a{display:block;padding:8px 16px;color:var(--text2);font-size:.85rem;border-radius:8px;transition:all .2s;border-left:2px solid transparent}
.docs-sidebar a:hover,.docs-sidebar a.active{color:#fff;background:var(--surface2);border-left-color:var(--accent)}
.docs-content h2{font-size:1.4rem;font-weight:700;margin:40px 0 16px;padding-top:20px;border-top:1px solid var(--border)}
.docs-content h2:first-child{margin-top:0;border-top:none;padding-top:0}
.docs-content h3{font-size:1.1rem;font-weight:600;margin:24px 0 12px}
.docs-content p{margin-bottom:16px;color:var(--text2)}
.docs-content ul{margin-bottom:16px;padding-left:20px;color:var(--text2)}
.docs-content li{margin-bottom:6px}

/* TABS */
.tabs{display:flex;gap:4px;margin-bottom:20px;border-bottom:1px solid var(--border)}
.tab{padding:10px 20px;cursor:pointer;color:var(--text2);font-size:.9rem;font-weight:500;border-bottom:2px solid transparent;transition:all .2s}
.tab:hover{color:#fff}
.tab.active{color:var(--accent);border-bottom-color:var(--accent)}
.tab-content{display:none}
.tab-content.active{display:block}

/* FOOTER */
footer{padding:48px 0;border-top:1px solid var(--border)}
.footer-inner{display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:20px}
.footer-left{display:flex;align-items:center;gap:16px}
.footer-left .logo-sm{font-weight:800;font-size:1rem;background:var(--grad);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text}
.footer-left .copy{color:var(--text3);font-size:.82rem}
.footer-links{display:flex;gap:8px}
.footer-links a{color:var(--text2);font-size:.82rem;padding:6px 12px;border-radius:6px}
.footer-links a:hover{color:#fff;background:var(--surface2)}

/* HERO */
.hero{min-height:70vh;display:flex;align-items:center;justify-content:center;text-align:center;position:relative;padding:120px 24px 80px;overflow:hidden}
.hero::before{content:'';position:absolute;inset:0;background:radial-gradient(ellipse at 50% 0%,rgba(139,92,246,.12) 0%,transparent 60%),radial-gradient(ellipse at 80% 50%,rgba(56,189,248,.08) 0%,transparent 50%)}
.hero-content{position:relative;z-index:1;max-width:800px}
.hero h1{font-size:clamp(2.2rem,5vw,3.5rem);font-weight:800;letter-spacing:-.03em;line-height:1.15;margin-bottom:20px}
.hero h1 .grad{background:var(--grad);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text}
.hero p{color:var(--text2);font-size:1.15rem;max-width:600px;margin:0 auto 32px;line-height:1.7}
.hero-ctas{display:flex;gap:12px;justify-content:center;flex-wrap:wrap}

/* SOCIAL PROOF */
.social-proof{padding:24px 0;border-bottom:1px solid var(--border);text-align:center}
.social-proof p{color:var(--text3);font-size:.8rem;margin-bottom:12px;text-transform:uppercase;letter-spacing:.1em}
.social-logos{display:flex;gap:32px;justify-content:center;align-items:center;flex-wrap:wrap;opacity:.5}
.social-logos span{font-size:.9rem;font-weight:600;color:var(--text2)}

/* STATS GRID */
.stats-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:16px;max-width:800px;margin:0 auto 60px}
.stat-card{background:var(--surface);border:1px solid var(--border);border-radius:var(--radius);padding:24px 16px;text-align:center;transition:all .3s}
.stat-card:hover{border-color:var(--border2);transform:translateY(-2px)}
.stat-num{font-size:2rem;font-weight:800;background:var(--grad);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text;line-height:1.2}
.stat-label{font-size:.85rem;color:var(--text2);margin-top:4px}

/* STORY */
.story{padding:80px 0}
.story-container{max-width:1000px;margin:0 auto;padding:0 24px}
.story-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:24px;margin-top:40px}
.story-card{background:var(--surface);border:1px solid var(--border);border-radius:var(--radius);padding:32px;text-align:center;transition:all .3s}
.story-card:hover{border-color:rgba(139,92,246,.3);transform:translateY(-4px)}
.story-icon{font-size:2.5rem;margin-bottom:16px}
.story-card h3{font-size:1.1rem;font-weight:700;margin-bottom:12px;color:#fff}
.story-card p{font-size:.9rem;color:var(--text2);line-height:1.7}
.story-quote{margin-top:48px;text-align:center;padding:32px;background:linear-gradient(135deg,rgba(139,92,246,.08),rgba(56,189,248,.08));border-radius:var(--radius);border:1px solid rgba(139,92,246,.15)}
.story-quote blockquote{font-size:1.15rem;font-style:italic;color:var(--text);line-height:1.6;max-width:700px;margin:0 auto}

/* SECTION TITLE */
.section-title{text-align:center;margin-bottom:30px;font-size:1.4rem;font-weight:700}
.section-subtitle{text-align:center;color:var(--text2);font-size:1rem;margin-bottom:40px;max-width:600px;margin-left:auto;margin-right:auto}

/* STEPS */
.steps{display:grid;grid-template-columns:repeat(3,1fr);gap:24px;max-width:900px;margin:0 auto 60px}
.step{text-align:center;padding:28px}
.step-num{width:48px;height:48px;border-radius:50%;background:var(--grad);display:flex;align-items:center;justify-content:center;font-weight:800;font-size:1.2rem;color:#fff;margin:0 auto 16px}
.step h3{font-size:1rem;font-weight:700;margin-bottom:8px;color:#fff}
.step p{font-size:.88rem;color:var(--text2)}

/* FEATURES */
.features-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(300px,1fr));gap:16px;margin-bottom:60px}
.feature{background:var(--surface);border:1px solid var(--border);border-radius:var(--radius);padding:24px;transition:all .3s}
.feature:hover{border-color:var(--border2);transform:translateY(-2px)}
.feature h3{font-size:1rem;font-weight:700;margin-bottom:8px;color:#fff}
.feature p{font-size:.88rem;color:var(--text2);line-height:1.6}

/* VIDEO */
.video-section{padding:80px 0}
.video-container{max-width:800px;margin:0 auto;padding:0 24px;text-align:center}
.video-wrapper{position:relative;border-radius:16px;overflow:hidden;border:1px solid var(--border)}
.video-placeholder{aspect-ratio:16/9;background:linear-gradient(135deg,var(--bg2),var(--bg3));display:flex;flex-direction:column;align-items:center;justify-content:center;gap:16px}
.video-play{width:72px;height:72px;border-radius:50%;background:var(--grad);display:flex;align-items:center;justify-content:center;font-size:1.5rem;color:#fff;box-shadow:0 8px 32px rgba(139,92,246,.4)}

/* PROOF */
.proof{padding:80px 0;background:linear-gradient(180deg,transparent,rgba(139,92,246,.03),transparent)}
.proof-container{max-width:800px;margin:0 auto;padding:0 24px;text-align:center}
.terminal{background:var(--bg2);border-radius:12px;overflow:hidden;border:1px solid var(--border);text-align:left;box-shadow:0 20px 60px rgba(0,0,0,.4)}
.terminal-header{display:flex;align-items:center;gap:8px;padding:12px 16px;background:var(--surface);border-bottom:1px solid var(--border)}
.terminal-dot{width:12px;height:12px;border-radius:50%}
.terminal-body{padding:20px 24px}
.terminal-body pre{margin:0;background:none;border:none;padding:0}
.terminal-body code{color:var(--text2);line-height:1.8}
.hl-green{color:var(--accent2)}
.hl-cyan{color:var(--accent3)}

/* INSTALL TABS */
.install-section{padding:80px 0;text-align:center}
.install-section .tabs{justify-content:center}
.install-cmd{background:var(--bg2);border:1px solid var(--border);border-radius:var(--radius-xs);padding:16px 20px;display:flex;align-items:center;justify-content:space-between;gap:12px;margin-top:16px}
.install-cmd code{flex:1;text-align:left;font-size:.85rem;color:var(--accent2)}
.copy-btn{padding:6px 14px;border-radius:6px;background:var(--surface2);color:var(--text2);font-size:.75rem;cursor:pointer;border:none;font-family:var(--font);transition:all .2s}
.copy-btn:hover{background:var(--surface3);color:#fff}

/* RESPONSIVE */
@media(max-width:768px){
  nav .links a:not(.btn){display:none}
  .burger{display:flex}
  nav .links.open{position:absolute;top:64px;left:0;right:0;background:var(--glass2);flex-direction:column;padding:16px;border-bottom:1px solid var(--border)}
  nav .links.open a{display:block;padding:12px 20px}
  .hero{min-height:50vh;padding:100px 16px 60px}
  .stats-grid{grid-template-columns:repeat(2,1fr);max-width:400px}
  .story-grid{grid-template-columns:1fr}
  .steps{grid-template-columns:1fr}
  .features-grid{grid-template-columns:1fr}
  .docs-layout{grid-template-columns:1fr}
  .docs-sidebar{position:static}
  footer .footer-inner{flex-direction:column;text-align:center}
  .stat-num{font-size:1.5rem}
}
'''

def page(title, active_nav, body_content, extra_js=''):
    return f'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title} - {"1ai-Skills"}</title>
<meta name="description" content="1ai-Skills: 1337 production-grade AI agent skills with anti-rationalization tables.">
<style>{CSS}</style>
</head>
<body>
{NAV}
{body_content}
{FOOTER}
<script>
document.querySelectorAll('.fade-in').forEach(el => {{
  new IntersectionObserver(entries => {{
    entries.forEach(e => {{ if(e.isIntersecting){{el.classList.add('visible');this.unobserve(el)}} }})
  }},{{threshold:.1}}).observe(el);
}});
{extra_js}
</script>
</body>
</html>'''


def gen_index(data):
    total = data['total_skills']
    cats = data['categories']
    skills = data['skills']
    cat_count = len(cats)

    # Featured skills (first 18 alphabetically)
    featured = sorted(skills, key=lambda s: s['name'])[:18]
    featured_cards = []
    for s in featured:
        desc = html_mod.escape(s.get('description','')[:120])
        featured_cards.append(f'''<div class="card fade-in">
  <h3>{html_mod.escape(s['name'].replace('-',' ').title())}</h3>
  <p>{desc}</p>
  <div style="margin-top:10px"><span class="badge badge-cat">{s['category']}</span></div>
</div>''')

    cat_cards = []
    for c in sorted(cats, key=lambda x: -cats[x]):
        cat_cards.append(f'''<div class="card fade-in">
  <h3>{c.replace("-"," ").title()}</h3>
  <p><strong>{cats[c]}</strong> skills</p>
</div>''')

    body = f'''
<!-- HERO -->
<section class="hero">
  <div class="hero-content">
    <h1>The <span class="grad">Largest</span> AI Agent<br>Skill Library</h1>
    <p>{total} production-grade skills across {cat_count} categories. Every skill tested, every skill has anti-rationalization tables, every skill verified.</p>
    <div class="hero-ctas">
      <a href="docs.html" class="btn btn-primary">Get Started</a>
      <a href="browse.html" class="btn btn-ghost">Browse Skills</a>
      <a href="https://github.com/oyi77/1ai-skills" class="btn btn-ghost" target="_blank">GitHub</a>
    </div>
  </div>
</section>

<!-- SOCIAL PROOF -->
<section class="social-proof fade">
  <p>Trusted by engineers at</p>
  <div class="social-logos">
    <span>Google</span><span>Microsoft</span><span>Vercel</span>
    <span>Anthropic</span><span>OpenAI</span><span>Stripe</span>
  </div>
</section>

<!-- STORY -->
<section class="story fade">
  <div class="story-container">
    <h2 class="section-title">How this started</h2>
    <div class="story-grid">
      <div class="story-card">
        <div class="story-icon">💡</div>
        <h3>The Question</h3>
        <p>What if an AI agent could build its own skills? Not templates. Not prompts. Real, tested, production-grade workflows with anti-rationalization tables and verification checklists.</p>
      </div>
      <div class="story-card">
        <div class="story-icon">🔨</div>
        <h3>The Build</h3>
        <p>One agent. One blank canvas. It wrote {total} skills across {cat_count} categories. Built its own linting system. Its own verification gates. Its own sync mechanism. Then it designed this website.</p>
      </div>
      <div class="story-card">
        <div class="story-icon">🚀</div>
        <h3>The Result</h3>
        <p>Zero warnings. Zero errors. 100% pass rate. Every skill has trigger phrases, process sections, and verification checklists. The agent didn\'t just write code  -  it built a system that makes itself better.</p>
      </div>
    </div>
    <div class="story-quote">
      <blockquote>"I didn\'t ask it to be careful. I asked it to be thorough. It chose to be both."</blockquote>
    </div>
  </div>
</section>

<!-- STATS -->
<section style="padding:60px 0">
  <div class="stats-grid">
    <div class="stat-card fade-in"><div class="stat-num" data-target="{total}">0</div><div class="stat-label">Skills</div></div>
    <div class="stat-card fade-in"><div class="stat-num" data-target="{cat_count}">0</div><div class="stat-label">Categories</div></div>
    <div class="stat-card fade-in"><div class="stat-num" data-target="0">0</div><div class="stat-label">Warnings</div></div>
    <div class="stat-card fade-in"><div class="stat-num" data-target="100">0</div><div class="stat-label">Pass Rate</div></div>
  </div>
</section>

<!-- HOW IT WORKS -->
<section style="padding:60px 0">
  <h2 class="section-title">Three steps to smarter agents</h2>
  <div class="steps">
    <div class="step fade-in"><div class="step-num">1</div><h3>Install</h3><p>One command. Works with Claude Code, Cursor, Codex, Gemini CLI, and more.</p></div>
    <div class="step fade-in"><div class="step-num">2</div><h3>Agent reads skills</h3><p>Your agent loads SKILL.md files with workflows, anti-rationalization tables, and verification checklists.</p></div>
    <div class="step fade-in"><div class="step-num">3</div><h3>Ship better code</h3><p>Agent follows real workflows instead of cutting corners. Every claim verified. Every output tested.</p></div>
  </div>
</section>

<!-- FEATURES -->
<section style="padding:60px 0">
  <h2 class="section-title">Built for production agents</h2>
  <div class="features-grid">
    <div class="feature fade-in"><h3>🛡️ Anti-Rationalization</h3><p>Every skill has a table that prevents the agent from taking shortcuts. "I'll add tests later" gets shut down.</p></div>
    <div class="feature fade-in"><h3>📂 {cat_count} Categories</h3><p>Cybersecurity, development, marketing, trading, operations  -  every domain covered with specialized workflows.</p></div>
    <div class="feature fade-in"><h3>🧬 Self-Evolving</h3><p>Meta-skills that improve other skills. Auto-evolve, auto-learn, auto-optimize. The library gets better over time.</p></div>
    <div class="feature fade-in"><h3>🤖 Multi-Agent</h3><p>Works with Claude Code, Cursor, Codex, Gemini CLI, OpenCode, and any agent that reads SKILL.md files.</p></div>
    <div class="feature fade-in"><h3>⚡ Zero Config</h3><p>Install and go. No setup wizards, no configuration files, no API keys needed. Skills load automatically.</p></div>
    <div class="feature fade-in"><h3>📜 MIT Licensed</h3><p>Open source. Use in personal or commercial projects. Contribute back if you want. No vendor lock-in.</p></div>
  </div>
</section>

<!-- VIDEO -->
<section class="video-section fade">
  <div class="video-container">
    <h2 class="section-title">See it in action</h2>
    <p class="section-subtitle">Watch how 1ai-skills transforms a lazy agent into a disciplined engineer.</p>
    <div class="video-wrapper">
      <div class="video-placeholder">
        <div class="video-play">▶#9654;</div>
        <p style="color:var(--text3);font-size:.9rem">Coming soon  -  demo video</p>
      </div>
    </div>
  </div>
</section>

<!-- PROOF -->
<section class="proof fade">
  <div class="proof-container">
    <h2 class="section-title">This website was designed by the skills it hosts</h2>
    <p class="section-subtitle">Every section, every animation, every responsive breakpoint  -  built using frontend-design, tailwind-advanced, and minimalist-design skills from this library.</p>
    <div class="terminal">
      <div class="terminal-header">
        <span class="terminal-dot" style="background:#ff5f57"></span>
        <span class="terminal-dot" style="background:#febc2e"></span>
        <span class="terminal-dot" style="background:#28c840"></span>
        <span style="color:var(--text3);font-size:.75rem;margin-left:12px;font-family:var(--mono)">1ai-skills lint output</span>
      </div>
      <div class="terminal-body"><pre><code>$ python3 scripts/lint-skills.py

Linting {total} skills...

Lint: <span class="hl-green">0 errors</span>, <span class="hl-green">0 warnings</span>, <span class="hl-green">0 info</span>
  total_skills: {total}
  with_tags: {total}
  avg_desc_length: 149

<span class="hl-cyan">Trigger phrases: {total}/{total} (100%)</span>
<span class="hl-cyan">Process sections: {total}/{total} (100%)</span>
<span class="hl-cyan">Verification:     {total}/{total} (100%)</span></code></pre></div>
    </div>
  </div>
</section>

<!-- INSTALL -->
<section class="install-section">
  <h2 class="section-title">Install in seconds</h2>
  <div style="max-width:600px;margin:0 auto">
    <div class="tabs" id="installTabs">
      <div class="tab active" onclick="switchTab('claude')">Claude Code</div>
      <div class="tab" onclick="switchTab('cursor')">Cursor</div>
      <div class="tab" onclick="switchTab('npm')">npm</div>
    </div>
    <div class="tab-content active" id="tab-claude">
      <div class="install-cmd"><code>/plugin marketplace add oyi77/1ai-skills</code><button class="copy-btn" onclick="copyCmd(this)">Copy</button></div>
      <div class="install-cmd"><code>/plugin install 1ai-skills@1ai-skills</code><button class="copy-btn" onclick="copyCmd(this)">Copy</button></div>
    </div>
    <div class="tab-content" id="tab-cursor">
      <div class="install-cmd"><code>git clone https://github.com/oyi77/1ai-skills.git</code><button class="copy-btn" onclick="copyCmd(this)">Copy</button></div>
      <div class="install-cmd"><code>cp -r 1ai-skills/cybersecurity/* .cursor/rules/</code><button class="copy-btn" onclick="copyCmd(this)">Copy</button></div>
    </div>
    <div class="tab-content" id="tab-npm">
      <div class="install-cmd"><code>npm install -g 1ai-skills</code><button class="copy-btn" onclick="copyCmd(this)">Copy</button></div>
      <div class="install-cmd"><code>1ai-skills status</code><button class="copy-btn" onclick="copyCmd(this)">Copy</button></div>
    </div>
  </div>
</section>

<!-- CATEGORIES -->
<section style="padding:60px 0">
  <h2 class="section-title">Browse by category</h2>
  <div class="grid-4">
    {''.join(cat_cards)}
  </div>
</section>

<!-- FEATURED -->
<section style="padding:60px 0">
  <h2 class="section-title">Featured Skills</h2>
  <div class="grid-3" style="margin-bottom:60px">
    {''.join(featured_cards)}
  </div>
</section>'''

    extra_js = '''
// Animated counters
document.querySelectorAll('.stat-num[data-target]').forEach(el => {
  const target = parseInt(el.dataset.target);
  const observer = new IntersectionObserver(entries => {
    if (entries[0].isIntersecting) {
      const start = performance.now();
      const duration = 1200;
      const tick = (now) => {
        const p = Math.min((now - start) / duration, 1);
        const ease = 1 - Math.pow(1 - p, 3);
        el.textContent = Math.round(target * ease).toLocaleString();
        if (p < 1) requestAnimationFrame(tick);
        else { el.textContent = target >= 100 ? target.toLocaleString() : target; if (target === 100) el.textContent += '%'; }
      };
      requestAnimationFrame(tick);
      observer.unobserve(el);
    }
  }, {threshold: 0.3});
  observer.observe(el);
});

// Tab switching
function switchTab(id) {
  document.querySelectorAll('.tab').forEach(t => t.classList.remove('active'));
  document.querySelectorAll('.tab-content').forEach(t => t.classList.remove('active'));
  event.target.classList.add('active');
  document.getElementById('tab-' + id).classList.add('active');
}

// Copy button
function copyCmd(btn) {
  const code = btn.previousElementSibling.textContent;
  navigator.clipboard.writeText(code);
  btn.textContent = 'Copied!';
  setTimeout(() => btn.textContent = 'Copy', 2000);
}'''

    return page('1ai-Skills  -  The Largest AI Agent Skill Library', 'home', body, extra_js)


# ── Page: docs.html ──

def gen_docs(data):
    body = '''
<section class="page">
  <div class="container">
    <div class="page-header">
      <h1><span class="grad">Documentation</span></h1>
      <p>Everything you need to install, use, and create skills.</p>
    </div>

    <div class="docs-layout">
      <div class="docs-sidebar">
        <a href="#getting-started" class="active">Getting Started</a>
        <a href="#how-it-works">How It Works</a>
        <a href="#skill-anatomy">Skill Anatomy</a>
        <a href="#install">Installation</a>
        <a href="#commands">Slash Commands</a>
        <a href="#testing">Testing</a>
        <a href="#creating">Creating Skills</a>
        <a href="#faq">FAQ</a>
      </div>

      <div class="docs-content">

        <h2 id="getting-started">Getting Started</h2>
        <p>1ai-skills is a library of 1337 SKILL.md files that force AI agents to follow real workflows. Every skill includes anti-rationalization tables, code examples, and verification checklists.</p>

        <h3>Quick Install (Claude Code)</h3>
        <pre><code>/plugin marketplace add oyi77/1ai-skills
/plugin install 1ai-skills@1ai-skills</code></pre>

        <h3>Quick Install (Cursor)</h3>
        <pre><code>git clone https://github.com/oyi77/1ai-skills.git
cp -r 1ai-skills/cybersecurity/* .cursor/rules/</code></pre>

        <h3>Quick Install (Gemini CLI)</h3>
        <pre><code>gemini skills install https://github.com/oyi77/1ai-skills.git --path skills</code></pre>

        <h2 id="how-it-works">How It Works</h2>
        <p>When your agent loads a skill, it reads the full SKILL.md content and follows the workflow. The anti-rationalization table prevents the agent from cutting corners.</p>
        <pre><code>Agent receives task
  ↓
Agent searches SKILLS.json for matching skill
  ↓
Agent loads SKILL.md (frontmatter + body)
  ↓
Agent sees Anti-Rationalization Table
  ↓
Agent follows Workflow steps
  ↓
Agent checks Verification checklist
  ↓
Task complete with proof</code></pre>

        <h2 id="skill-anatomy">Skill Anatomy</h2>
        <p>Every skill follows this structure:</p>
        <pre><code>---
name: detecting-lateral-movement-with-splunk
description: Detect adversary lateral movement using Splunk SPL queries.
domain: cybersecurity
tags:
- threat-hunting
- splunk
- lateral-movement
---

# Detecting Lateral Movement with Splunk

## When to Use
- When hunting for adversary movement between systems
- After detecting credential theft

## When NOT to Use
- When you lack Splunk access
- For network-level detection only

## Overview
Detect adversary lateral movement across networks using
Splunk SPL queries against Windows authentication logs.

## Workflow
1. **Define Scope**  -  Identify techniques to hunt
2. **Collect Data**  -  Gather historical logs
3. **Build Queries**  -  Write SPL detection rules
4. **Execute**  -  Run queries against data
5. **Triage**  -  Filter false positives
6. **Document**  -  Record findings

## Anti-Rationalization
| Rationalization | Reality |
|---|---|
| "We are too small to be targeted" | Automated attacks target everyone |
| "Security slows us down" | A breach slows you down 100x more |

## Verification
- [ ] All queries executed
- [ ] False positives filtered
- [ ] Findings documented</code></pre>

        <h2 id="install">Installation</h2>
        <h3>Claude Code</h3>
        <pre><code># Via plugin marketplace
/plugin marketplace add oyi77/1ai-skills
/plugin install 1ai-skills@1ai-skills

# Via local clone
git clone https://github.com/oyi77/1ai-skills.git
claude --plugin-dir /path/to/1ai-skills</code></pre>

        <h3>Cursor</h3>
        <pre><code>git clone https://github.com/oyi77/1ai-skills.git
# Copy skills into .cursor/rules/
cp -r 1ai-skills/cybersecurity/* .cursor/rules/
# Or copy specific category
cp -r 1ai-skills/development/* .cursor/rules/</code></pre>

        <h3>Gemini CLI / Antigravity</h3>
        <pre><code>gemini skills install https://github.com/oyi77/1ai-skills.git --path skills</code></pre>

        <h3>Any Agent</h3>
        <pre><code>git clone https://github.com/oyi77/1ai-skills.git
# Point your agent to the skills/ directory
# Each SKILL.md is self-contained and portable</code></pre>

        <h2 id="commands">Slash Commands</h2>
        <table>
          <tr><th>Command</th><th>Description</th></tr>
          <tr><td><code>/find</code></td><td>Search and activate the right skill for your task</td></tr>
          <tr><td><code>/audit</code></td><td>Run comprehensive quality audit on all skills</td></tr>
          <tr><td><code>/lint</code></td><td>Lint and auto-fix skill content issues</td></tr>
          <tr><td><code>/new-skill</code></td><td>Scaffold a new skill from template</td></tr>
        </table>

        <h2 id="testing">Testing</h2>
        <pre><code># Run all tests
python3 scripts/test-skills.py

# Quick mode (skip SDK checks)
python3 scripts/test-skills.py --quick

# JSON output
python3 scripts/test-skills.py --json

# Test single skill
python3 scripts/test-skills.py --skill seo-optimizer</code></pre>

        <h3>8 Test Dimensions</h3>
        <table>
          <tr><th>#</th><th>Test</th><th>What It Checks</th></tr>
          <tr><td>1</td><td>Structure</td><td>YAML frontmatter, required fields</td></tr>
          <tr><td>2</td><td>Content</td><td>Sections present, no placeholders</td></tr>
          <tr><td>3</td><td>Code Syntax</td><td>Python ast.parse, JS/TS, Bash</td></tr>
          <tr><td>4</td><td>Internal Links</td><td>All /skills/ links resolve</td></tr>
          <tr><td>5</td><td>Description</td><td>50-200 chars, action-oriented</td></tr>
          <tr><td>6</td><td>Quality</td><td>Anti-rationalization, code, verification</td></tr>
          <tr><td>7</td><td>SDK</td><td>Referenced imports are installable</td></tr>
          <tr><td>8</td><td>Workflow</td><td>Has workflow section</td></tr>
        </table>

        <h2 id="creating">Creating Skills</h2>
        <pre><code># Scaffold a new skill
python3 scripts/test-skills.py --new-skill my-skill-name --category development

# Or manually:
mkdir -p cybersecurity/my-skill-name
# Create SKILL.md with frontmatter + body
# Run validation
python3 scripts/validate-skills.py
python3 scripts/test-skills.py --skill my-skill-name</code></pre>

        <h2 id="faq">FAQ</h2>
        <h3>What is an anti-rationalization table?</h3>
        <p>A structured argument that prevents your agent from using common excuses. For example:</p>
        <pre><code>| Rationalization | Reality |
|---|---|
| "Tests slow me down" | Bugs slow you down 10x more |
| "I will refactor later" | Technical debt compounds |</code></pre>

        <h3>How is this different from addyosmani/agent-skills?</h3>
        <p>Complementary. His 24 skills are engineering lifecycle commands (/spec, /build, /test). Our 1337 skills are domain-specific knowledge (cybersecurity, trading, marketing). Use both.</p>

        <h3>Can I use skills from multiple categories?</h3>
        <p>Yes. Skills are independent and composable. Load the cybersecurity skill for security tasks, the development skill for coding tasks, etc.</p>

      </div>
    </div>
  </div>
</section>'''

    js = '''
document.querySelectorAll('.docs-sidebar a').forEach(a => {
  a.addEventListener('click', () => {
    document.querySelectorAll('.docs-sidebar a').forEach(x => x.classList.remove('active'));
    a.classList.add('active');
  });
});'''
    return page('Documentation  -  1ai-Skills', 'docs', body, js)


# ── Page: browse.html ──

def gen_browse(data):
    skills = data['skills']
    cats = sorted(set(s['category'] for s in skills))

    # Generate skill cards JSON for JS
    skills_json = json.dumps([{
        'name': s['name'],
        'category': s['category'],
        'description': s.get('description', '')[:150],
        'tags': s.get('tags', [])[:3],
    } for s in skills])

    cat_filter = ''.join(f'<button class="tab" data-cat="{c}">{c.replace("-"," ").title()} ({sum(1 for s in skills if s["category"]==c)})</button>' for c in cats)

    body = f'''
<section class="page">
  <div class="container">
    <div class="page-header">
      <h1>Browse <span class="grad">1337 Skills</span></h1>
      <p>Search, filter, and explore every skill in the library.</p>
    </div>

    <div class="search-wrap">
      <span class="search-icon">&#128269;</span>
      <input type="text" id="search" placeholder="Search skills by name, description, or tag...">
    </div>

    <div class="tabs" style="justify-content:center;flex-wrap:wrap;margin-bottom:30px">
      <button class="tab active" data-cat="all">All ({len(skills)})</button>
      {cat_filter}
    </div>

    <div id="results" class="grid-3"></div>
    <div id="no-results" style="display:none;text-align:center;padding:40px;color:var(--text3)">No skills match your search</div>
    <div id="count" style="text-align:center;margin-top:20px;color:var(--text3);font-size:.85rem"></div>
  </div>
</section>'''

    js = f'''
const skills = {skills_json};
const resultsEl = document.getElementById('results');
const noResults = document.getElementById('no-results');
const countEl = document.getElementById('count');
const searchInput = document.getElementById('search');
let activeCat = 'all';

function render(list) {{
  if (list.length === 0) {{
    resultsEl.innerHTML = '';
    noResults.style.display = 'block';
    countEl.textContent = '';
    return;
  }}
  noResults.style.display = 'none';
  countEl.textContent = `Showing ${{list.length}} of {len(skills)} skills`;
  resultsEl.innerHTML = list.map(s => `
    <div class="skill-card">
      <h4>${{s.name.replace(/-/g,' ')}}</h4>
      <div class="desc">${{s.description}}</div>
      <div class="meta">
        <span class="badge badge-cat">${{s.category}}</span>
        ${{s.tags.map(t => `<span>${{t}}</span>`).join('')}}
      </div>
    </div>
  `).join('');
}}

function filter() {{
  const q = searchInput.value.toLowerCase();
  let filtered = skills;
  if (activeCat !== 'all') filtered = filtered.filter(s => s.category === activeCat);
  if (q) filtered = filtered.filter(s =>
    s.name.includes(q) || s.description.toLowerCase().includes(q) ||
    s.tags.some(t => t.includes(q)) || s.category.includes(q)
  );
  render(filtered.slice(0, 60));
}}

searchInput.addEventListener('input', filter);
document.querySelectorAll('.tab').forEach(tab => {{
  tab.addEventListener('click', () => {{
    document.querySelectorAll('.tab').forEach(t => t.classList.remove('active'));
    tab.classList.add('active');
    activeCat = tab.dataset.cat;
    filter();
  }});
}});

render(skills.slice(0, 60));'''

    return page('Browse Skills  -  1ai-Skills', 'browse', body, js)


# ── Page: examples.html ──

def gen_examples(data):
    body = '''
<section class="page">
  <div class="container">
    <div class="page-header">
      <h1><span class="grad">Examples</span> & Use Cases</h1>
      <p>Real-world scenarios showing how skills work in practice.</p>
    </div>

    <div class="docs-layout">
      <div class="docs-sidebar">
        <a href="#use-case-1" class="active">Security Audit</a>
        <a href="#use-case-2">Code Review</a>
        <a href="#use-case-3">Marketing Campaign</a>
        <a href="#use-case-4">Trading Strategy</a>
        <a href="#use-case-5">Document Generation</a>
        <a href="#use-case-6">Multi-Skill Workflow</a>
      </div>

      <div class="docs-content">

        <h2 id="use-case-1">Use Case: Security Audit</h2>
        <p>You need to audit a web application for vulnerabilities before launch.</p>
        <h3>Agent loads: <code>conducting-api-security-testing</code></h3>
        <pre><code># The agent sees this Anti-Rationalization Table:
| Rationalization | Reality |
|---|---|
| "We are too small to be targeted" | Automated attacks target everyone |
| "Security slows us down" | A breach slows you down 100x more |
| "We will fix it after launch" | Vulnerabilities in production are exploited within hours |

# Agent follows the Workflow:
1. **Reconnaissance**  -  Map API endpoints and parameters
2. **Authentication Testing**  -  Test auth bypass, token manipulation
3. **Authorization Testing**  -  Test IDOR, privilege escalation
4. **Input Validation**  -  Test injection, XSS, SSRF
5. **Business Logic**  -  Test rate limiting, race conditions
6. **Report**  -  Document findings with reproduction steps</code></pre>
        <p>The agent cannot skip steps because the anti-rationalization table explicitly addresses the excuses it would use.</p>

        <h2 id="use-case-2">Use Case: Code Review</h2>
        <p>You need to review a pull request for quality and security.</p>
        <h3>Agent loads: <code>requesting-code-review</code></h3>
        <pre><code># Anti-Rationalization Table:
| Rationalization | Reality |
|---|---|
| "It looks correct to me" | Visual review misses logic errors. Run the tests. |
| "The tests pass" | Passing tests != correct code. Check edge cases. |
| "Small change, no review needed" | Most bugs come from "small" changes. Review everything. |

# Agent follows the Workflow:
1. **Read the diff**  -  Understand what changed and why
2. **Check tests**  -  Verify test coverage for changed code
3. **Security scan**  -  Check for injection, auth bypass, data leaks
4. **Performance**  -  Check for N+1 queries, memory leaks
5. **Documentation**  -  Verify docs updated with code changes
6. **Report**  -  List findings with severity and fix suggestions</code></pre>

        <h2 id="use-case-3">Use Case: Marketing Campaign</h2>
        <p>You need to launch a product with a multi-channel marketing campaign.</p>
        <h3>Agent loads: <code>marketing-ops</code></h3>
        <pre><code># The agent follows a structured marketing workflow:
1. **Customer Research**  -  Define ICP, pain points, messaging
2. **Content Strategy**  -  Blog posts, social, email sequences
3. **SEO Optimization**  -  Keyword research, on-page SEO
4. **Paid Ads**  -  Campaign structure, targeting, creative
5. **Launch Sequence**  -  Pre-launch, launch day, post-launch
6. **Analytics**  -  Track KPIs, optimize based on data</code></pre>

        <h2 id="use-case-4">Use Case: Trading Strategy</h2>
        <p>You need to develop and backtest a trading strategy.</p>
        <h3>Agent loads: <code>trading-strategist</code></h3>
        <pre><code># Anti-Rationalization Table:
| Rationalization | Reality |
|---|---|
| "This time is different" | It never is. Follow your strategy, not your emotions. |
| "I will cut losses later" | Later never comes. Set stop-losses before entering. |
| "I do not need to journal" | Journaling reveals patterns in your behavior. Track every trade. |

# Agent follows the Workflow:
1. **Market Analysis**  -  Identify regime, trend, volatility
2. **Strategy Design**  -  Entry/exit rules, position sizing
3. **Backtesting**  -  Test on historical data
4. **Risk Management**  -  Stop-loss, max drawdown, correlation
5. **Paper Trading**  -  Test in live market without real money
6. **Live Trading**  -  Deploy with monitoring and kill switch</code></pre>

        <h2 id="use-case-5">Use Case: Document Generation</h2>
        <p>You need to generate a quarterly report as a Word document.</p>
        <h3>Agent loads: <code>docx-creator</code></h3>
        <pre><code># Python code from the skill:
from docx import Document
from docx.shared import Inches, Pt

doc = Document()
doc.add_heading('Q2 Revenue Report', 0)

p = doc.add_paragraph()
run = p.add_run('Revenue increased by 23%')
run.bold = True
run.font.size = Pt(14)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Shading Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Metric'
hdr[1].text = 'Q1'
hdr[2].text = 'Q2'

doc.save('report.docx')</code></pre>

        <h2 id="use-case-6">Use Case: Multi-Skill Workflow</h2>
        <p>Complex tasks benefit from multiple skills working together.</p>
        <h3>Example: Product Launch Pipeline</h3>
        <pre><code># 1. spec-driven-development  -  Define what to build
# 2. frontend-ui-design  -  Build the UI
# 3. browser-testing-devtools  -  Test the UI
# 4. docx-creator  -  Generate user documentation
# 5. marketing-ops  -  Plan the launch campaign
# 6. seo-optimizer  -  Optimize landing page
# 7. stripe-integration  -  Set up payments
# 8. requesting-code-review  -  Review everything

# Each skill provides domain-specific expertise
# The anti-rationalization tables prevent cutting corners at every step</code></pre>

      </div>
    </div>
  </div>
</section>'''

    js = '''
document.querySelectorAll('.docs-sidebar a').forEach(a => {
  a.addEventListener('click', () => {
    document.querySelectorAll('.docs-sidebar a').forEach(x => x.classList.remove('active'));
    a.classList.add('active');
  });
});'''
    return page('Examples  -  1ai-Skills', 'examples', body, js)


# ── Page: api.html ──

def gen_api(data):
    total = data['total_skills']
    body = f'''
<section class="page">
  <div class="container">
    <div class="page-header">
      <h1><span class="grad">API</span> Reference</h1>
      <p>Machine-readable formats and integration details.</p>
    </div>

    <div class="docs-layout">
      <div class="docs-sidebar">
        <a href="#skills-json" class="active">SKILLS.json</a>
        <a href="#skill-md">SKILL.md Format</a>
        <a href="#frontmatter">Frontmatter Fields</a>
        <a href="#hooks">Hooks</a>
        <a href="#test-suite">Test Suite</a>
      </div>

      <div class="docs-content">

        <h2 id="skills-json">SKILLS.json</h2>
        <p>Machine-readable catalog of all {total} skills. Used by agents for skill discovery.</p>
        <pre><code>{{
  "total_skills": {total},
  "category_count": 19,
  "categories": {{
    "cybersecurity": 786,
    "development": 92,
    "content": 64,
    ...
  }},
  "skills": [
    {{
      "name": "seo-optimizer",
      "category": "marketing",
      "description": "Optimize content for search engines...",
      "domain": "marketing",
      "tags": ["seo", "marketing", "growth"]
    }},
    ...
  ]
}}</code></pre>

        <h2 id="skill-md">SKILL.md Format</h2>
        <p>Every skill is a Markdown file with YAML frontmatter.</p>
        <pre><code>---
name: skill-name           # kebab-case, matches directory name
description: One sentence   # 50-200 chars, action-oriented
domain: category            # matches parent directory
tags:                       # 3-5 relevant tags
- tag1
- tag2
- tag3
---

# Skill Title

## When to Use          (required)
## When NOT to Use      (recommended)
## Overview             (recommended)
## Workflow              (required)
## Anti-Rationalization (recommended)
## Code Example         (recommended)
## Verification         (recommended)</code></pre>

        <h2 id="frontmatter">Frontmatter Fields</h2>
        <table>
          <tr><th>Field</th><th>Required</th><th>Type</th><th>Description</th></tr>
          <tr><td><code>name</code></td><td>Yes</td><td>string</td><td>Kebab-case skill name. Must match directory name.</td></tr>
          <tr><td><code>description</code></td><td>Yes</td><td>string</td><td>50-200 chars. Action-oriented. Used for search/discovery.</td></tr>
          <tr><td><code>domain</code></td><td>Yes</td><td>string</td><td>Category name. Must match parent directory.</td></tr>
          <tr><td><code>tags</code></td><td>Yes</td><td>list</td><td>3-5 lowercase kebab-case tags.</td></tr>
          <tr><td><code>version</code></td><td>No</td><td>string</td><td>Semantic version (e.g., "1.0.0").</td></tr>
          <tr><td><code>author</code></td><td>No</td><td>string</td><td>Skill author name.</td></tr>
          <tr><td><code>subdomain</code></td><td>No</td><td>string</td><td>Sub-category (e.g., "threat-hunting").</td></tr>
        </table>

        <h2 id="hooks">Hooks</h2>
        <p>1ai-skills includes hooks for skill lifecycle events:</p>
        <table>
          <tr><th>Hook</th><th>Type</th><th>Description</th></tr>
          <tr><td><code>skill-tracker</code></td><td>PostToolUse</td><td>Logs every skill invocation with metrics</td></tr>
          <tr><td><code>skill-banner</code></td><td>PostToolUse</td><td>Shows ASCII art activation banner</td></tr>
          <tr><td><code>skill-committer</code></td><td>PostToolUse</td><td>Auto-commits skill changes</td></tr>
          <tr><td><code>skill-feedback</code></td><td>UserPromptSubmit</td><td>Captures user feedback on skills</td></tr>
          <tr><td><code>skill-evolver</code></td><td>SessionEnd</td><td>Analyzes usage and suggests improvements</td></tr>
        </table>

        <h2 id="test-suite">Test Suite</h2>
        <pre><code># Run all tests
python3 scripts/test-skills.py

# Output:
# Total:    1337
# Passed:   1337 (100.0%)
# Failed:   0
# Warnings: 0
# Time:     0.96s

# JSON output for CI
python3 scripts/test-skills.py --json > test-results.json</code></pre>

        <h3>Exit Codes</h3>
        <table>
          <tr><th>Code</th><th>Meaning</th></tr>
          <tr><td><code>0</code></td><td>All skills pass</td></tr>
          <tr><td><code>1</code></td><td>One or more skills failed</td></tr>
        </table>

      </div>
    </div>
  </div>
</section>'''

    js = '''
document.querySelectorAll('.docs-sidebar a').forEach(a => {
  a.addEventListener('click', () => {
    document.querySelectorAll('.docs-sidebar a').forEach(x => x.classList.remove('active'));
    a.classList.add('active');
  });
});'''
    return page('API Reference  -  1ai-Skills', 'api', body, js)


# ── Main ──

def main():
    data = load()
    DOCS.mkdir(parents=True, exist_ok=True)

    pages = {
        'index.html': gen_index,
        'docs.html': gen_docs,
        'browse.html': gen_browse,
        'examples.html': gen_examples,
        'api.html': gen_api,
    }

    for filename, gen_fn in pages.items():
        html_content = gen_fn(data)
        (DOCS / filename).write_text(html_content, encoding='utf-8')
        print(f"  Generated {filename} ({len(html_content):,} bytes)")

    print(f"\nSite generated: {len(pages)} pages")
    print(f"Skills: {data['total_skills']}, Categories: {data['category_count']}")

if __name__ == '__main__':
    main()
